"""Build the RL Assignment-2 report as a .docx using the user's existing homework
report as the template (so it inherits the exact same styles, fonts and footer),
with shorter, simpler wording. Then this docx is converted to PDF with LibreOffice.
"""
import os
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = '/Users/Joachim/Downloads/Machine Learning for Artificial Intelligence Engineering - Homework 2 Report.docx'
OUT = os.path.join(HERE, 'RL26_Ass2_Report.docx')

# ---- DQN numbers from the real training run ----
d = np.load(os.path.join(HERE, 'dqn_metrics.npz'))
r = d['reward'].astype(float)
NE, FR, FEPS = len(r), int(d['current_step']), float(d['eps'])
R_FIRST, R_LAST, R_MAX = np.mean(r[:50]), np.mean(r[-50:]), int(r.max())

doc = Document(TEMPLATE)

# ---- clear the template body but keep the section properties (page size + footer).
# Keep the sectPr in place; python-docx inserts new content before it. ----
body = doc.element.body
sectPr = body.find(qn('w:sectPr'))
# the template (a Google-Docs export) stores page size/margins as floats like
# "1440.0000000000002"; python-docx needs integer twips, so round them.
for tag in ('w:pgSz', 'w:pgMar'):
    el = sectPr.find(qn(tag))
    if el is not None:
        for a, v in list(el.attrib.items()):
            try:
                el.set(a, str(int(round(float(v)))))
            except (ValueError, TypeError):
                pass
for child in list(body):
    if child is not sectPr:
        body.remove(child)


def style_para(text, style):
    p = doc.add_paragraph(style=style)
    if text:
        p.add_run(text)
    return p


def h1(t):
    return style_para(t, 'Heading 1')


def h2(t):
    return style_para(t, 'Heading 2')


def body_p(*segments, justify=True):
    """segments: strings, or (text, {'bold':True}) tuples."""
    p = doc.add_paragraph(style='Normal')
    if justify:
        p.alignment = ALIGN.JUSTIFY
    for seg in segments:
        if isinstance(seg, tuple):
            txt, opt = seg
            run = p.add_run(txt)
            run.bold = opt.get('bold', False)
            run.italic = opt.get('italic', False)
        else:
            p.add_run(seg)
    return p


def equation(text):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def caption(text):
    p = doc.add_paragraph(style='Normal')
    p.alignment = ALIGN.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.space_after = Pt(10)
    return p


def figure(fname, cap, width_in=6.1):
    path = os.path.join(HERE, fname)
    doc.add_picture(path, width=Inches(width_in))
    doc.paragraphs[-1].alignment = ALIGN.CENTER
    caption(cap)


def _set_borders(table, color='AAAAAA', sz=4):
    tblPr = table._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        b.append(e)
    tblPr.append(b)


def _shade(cell, fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def make_table(rows, widths, header=True, font=8.5):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = 1  # center
    _set_borders(table)
    for ci, w in enumerate(widths):
        for ri in range(len(rows)):
            table.cell(ri, ci).width = Cm(w)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.cell(ri, ci)
            cell.paragraphs[0].text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(font)
            if (header and ri == 0):
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if ci > 0 or (header and ri == 0):
                cell.paragraphs[0].alignment = ALIGN.CENTER
            if header and ri == 0:
                _shade(cell, '1D4E79')
            elif ri % 2 == 0:
                _shade(cell, 'EEF3F9')
    doc.add_paragraph(style='Normal').paragraph_format.space_after = Pt(2)
    return table


# =====================================================================
# TITLE + SUBMISSION DETAILS
# =====================================================================
style_para('Reinforcement Learning VU (708.006), Summer Term 2026, Assignment 2 Report', 'Title')

h1('Submission Details')
sub = make_table([
    ['Item', 'Value'],
    ['Student Name', 'Joachim Rath'],
    ['Student ID', '51811071'],
    ['Course', 'Reinforcement Learning VU (708.006)'],
    ['Due Date', '19.06.2026 23:59'],
], widths=[4.0, 11.5], header=True, font=10)

# =====================================================================
# PART 1
# =====================================================================
h1('1 Reinforcement Learning in a Grid World')
body_p("In this part the agent has to control the 4×4 slippery FrozenLake without knowing the transition "
       "model. It keeps a tabular Q(s,a) and learns by temporal-difference (TD) control while acting "
       "ε-greedily. I train each agent for 10 000 episodes with ε decaying from 1, and then test the greedy "
       "policy (ε = 0) for 5 000 episodes. The only reward is 1 for reaching the goal.")

h2('1.1 Algorithms')
body_p(("ε-greedy", {'bold': True}), " means a random action with probability ε and arg max over a of "
       "Q(s,a) otherwise; for testing I set ε = 0. The three methods only differ in the TD target used in "
       "Q(s,a) ← Q(s,a) + α·[target − Q(s,a)]. When the next state is terminal the bootstrap term is dropped:")
equation("SARSA:           target = r + γ · Q(s′, a′),   a′ from the policy")
equation("Q-Learning:      target = r + γ · max_a′ Q(s′, a′)")
equation("Expected SARSA:  target = r + γ · Σ_a′ π(a′|s′) · Q(s′, a′)")
body_p("For an ε-greedy policy the Expected-SARSA expectation is just "
       "ε · mean_a Q(s′,a) + (1−ε) · max_a Q(s′,a).")

h2('1.2 Hyperparameters and results')
body_p("I tuned α and eps_decay separately for each algorithm and γ with a grid search over "
       "α ∈ {0.01, 0.05, 0.1, 0.2, 0.5} and eps_decay ∈ {0.998, 0.999, 0.9993, 0.9995, 0.9997}, scoring "
       "each setting by the mean test return over 3 seeds. α = 0.01 was always too slow. For γ = 0.95 all "
       "three methods reached about 0.78 over a wide range of α. For γ = 1.0 Q-Learning was unstable: with "
       "α ≥ 0.05 and slow ε-decay it often diverged to a useless policy (return 0), and only α = 0.05 "
       "worked, while SARSA and Expected SARSA stayed stable. Table 1 shows the values I used and the "
       "returns I got; all six are close to Policy Iteration (0.786 for γ = 0.95, 0.826 for γ = 1.0).")
make_table([
    ['Algorithm', 'γ', 'α', 'eps_decay', 'train R', 'test R', 'PI test'],
    ['SARSA', '0.95', '0.05', '0.9993', '0.49', '0.77', '0.786'],
    ['Q-Learning', '0.95', '0.10', '0.9990', '0.55', '0.78', '0.786'],
    ['Expected SARSA', '0.95', '0.10', '0.9993', '0.48', '0.78', '0.786'],
    ['SARSA', '1.0', '0.10', '0.9997', '0.23', '0.82', '0.826'],
    ['Q-Learning', '1.0', '0.05', '0.9993', '0.51', '0.77', '0.826'],
    ['Expected SARSA', '1.0', '0.10', '0.9997', '0.24', '0.82', '0.826'],
], widths=[3.4, 1.2, 1.2, 2.2, 2.0, 1.8, 1.9], header=True)
caption('Table 1: Chosen hyperparameters (seed 0) and resulting mean training / test return, with the '
        'Policy-Iteration reference.')
body_p("Figures 1–6 show the training and test reward, the learned policy (white arrows) and the value "
       "function for each case.")
for fn, cap in [
    ('frozenlake_SARSA_gamma_0.95.png', 'Figure 1: SARSA, γ = 0.95.'),
    ('frozenlake_Q-Learning_gamma_0.95.png', 'Figure 2: Q-Learning, γ = 0.95.'),
    ('frozenlake_Expected_SARSA_gamma_0.95.png', 'Figure 3: Expected SARSA, γ = 0.95.'),
    ('frozenlake_SARSA_gamma_1.0.png', 'Figure 4: SARSA, γ = 1.0.'),
    ('frozenlake_Q-Learning_gamma_1.0.png', 'Figure 5: Q-Learning, γ = 1.0.'),
    ('frozenlake_Expected_SARSA_gamma_1.0.png', 'Figure 6: Expected SARSA, γ = 1.0.'),
]:
    figure(fn, cap)

h2('1.3 Comparison with Policy Iteration')
body_p("All three methods find essentially the same greedy policy as Policy Iteration and reach a similar "
       "success rate. The typical slippery trick is visible: next to a hole the policy points away from it, "
       "often into a wall, so a random sideways slip cannot fall in. The value functions agree well at "
       "γ = 0.95. At γ = 1.0 they differ (Table 2): Q-Learning still matches v* because it learns greedy "
       "values off-policy, while SARSA and Expected SARSA underestimate the states far from the goal. They "
       "are on-policy and evaluate the exploring ε-greedy policy, and without discounting the value spreads "
       "only slowly from the single goal reward, so far-away cells stay too low after 10 000 episodes.")
make_table([
    ['v(s) at γ = 1.0', 'start (0,0)', 'left-col (2,0)', 'left-of-goal (3,2)'],
    ['Policy Iteration (v*)', '0.824', '0.824', '0.941'],
    ['Q-Learning', '0.797', '0.788', '0.937'],
    ['SARSA', '0.504', '0.574', '0.900'],
    ['Expected SARSA', '0.514', '0.540', '0.824'],
], widths=[4.6, 3.2, 3.4, 3.6], header=True)
caption('Table 2: Value of three cells at γ = 1.0. Near the goal all agree; far from it only Q-Learning '
        'stays close to v*.')

# =====================================================================
# PART 2
# =====================================================================
h1('2 Deep Q-Learning for Atari Breakout')
body_p("Here I train a DQN (Mnih et al., 2015) on Atari Breakout with a convolutional network, a fixed "
       "target network and an experience-replay buffer. Frames are turned to grayscale, resized to 84×84 "
       "and the last 4 are stacked, so a state is a 4×84×84 tensor; there are 4 actions.")

h2('2.1 Network and fixed target network')
body_p("The online network has three convolutional layers (32 filters 8×8 stride 4, 64 filters 4×4 "
       "stride 2, 64 filters 3×3 stride 1, each with ReLU), then a 512-unit dense layer and a final layer "
       "with one output per action (about 1.7M parameters). A second network with the same shape is the "
       "target network. Its weights are frozen and only copied from the online network every sync_target "
       "frames. The target value is y = r + γ · max_a′ Q_target(s′,a′), and I minimise the MSE to "
       "Q_online(s,a) with Adam. The target network keeps this regression target fixed for a while instead "
       "of letting it move on every update, which stops the training from oscillating. The "
       "burn_in_phase is the number of frames at the start where the agent only fills the buffer (ε = 1, no "
       "updates, no ε-decay), so the first minibatches are not drawn from an almost empty, highly "
       "correlated buffer.")

h2('2.2 Experience replay memory')
body_p("The buffer is a deque of past transitions and sample() takes a uniform random minibatch. The "
       "reason for it is that SGD assumes the samples are i.i.d. (independent and identically distributed), "
       "which is violated if you learn from transitions in the order they happen: consecutive transitions "
       "are highly correlated, and the data distribution keeps changing as the policy changes "
       "(non-stationary). Sampling randomly from a large buffer mixes many episodes and policies, so the "
       "minibatches are roughly independent again, and every transition can be reused several times.")

h2('2.3 Results')
body_p(f"I trained on an Apple-M2 laptop (8 GB, MPS backend), which is far too little for Breakout (the "
       f"original DQN uses about 50M frames), so this is only a short demonstration. In about 40 minutes "
       f"the agent ran {NE} episodes ({FR:,} steps) and ε reached {FEPS:.1f}. The mean episode reward went "
       f"up from {R_FIRST:.0f} to {R_LAST:.0f} (best episode {R_MAX}), so it clearly learns to play better "
       f"than random, which scores about 2, and the loss stays stable after the burn-in (Figure 7). With a "
       f"GPU and a larger buffer the same code would reach higher scores.")
figure('dqn_train_curves.png', 'Figure 7: DQN on Breakout — reward per episode (left) and TD loss (right), '
       'each with a running mean.', width_in=6.3)

# =====================================================================
# LLM USAGE
# =====================================================================
h1('3 Use of LLMs')
body_p("I used Claude (Anthropic) through the Claude Code assistant. It helped me write the TD updates and "
       "the ε-greedy policy, run and interpret the hyperparameter search (including why Q-Learning is "
       "unstable at γ = 1.0), fill in the DQN skeleton (network, target sync, replay sampling, optimiser) "
       "and adapt the code to my 8 GB laptop. It also noticed that the agent was being created without the "
       "run_as_ddqn argument it expects, and that the per-episode loss divides by the step counter. The main "
       "limitations were that the given skeleton needed those fixes to run, and that I could not train "
       "Breakout to a high score on my hardware, so the DQN part stays a short demonstration. I went through "
       "the code and the reasoning myself and understand how it works.")

doc.save(OUT)
print('Wrote', OUT)
# sanity: section properties (page size + footer) still present and last
assert body.find(qn('w:sectPr')) is not None
assert list(body)[-1].tag == qn('w:sectPr'), 'sectPr is not the last body child!'
print('sectPr preserved as last child:', list(body)[-1].tag == qn('w:sectPr'))
